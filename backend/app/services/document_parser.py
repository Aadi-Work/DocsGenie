from __future__ import annotations

import io
import re
from pathlib import Path
from typing import Optional


SUPPORTED_EXTENSIONS = {".txt", ".md", ".docx", ".pdf", ".xlsx", ".csv"}


class DocumentParseError(Exception):
    pass


def extract_text_from_bytes(filename: str, data: bytes, max_chars: int = 20000) -> str:
    suffix = Path(filename).suffix.lower()
    if suffix not in SUPPORTED_EXTENSIONS:
        raise DocumentParseError(
            f"Unsupported file type '{suffix}'. Use: {', '.join(sorted(SUPPORTED_EXTENSIONS))}"
        )
    if suffix in {".txt", ".md", ".csv"}:
        text = _decode_text(data)
    elif suffix == ".docx":
        text = _extract_docx(data)
    elif suffix == ".pdf":
        text = _extract_pdf(data)
    elif suffix == ".xlsx":
        text = _extract_xlsx(data)
    else:
        text = _decode_text(data)

    text = re.sub(r"\n{3,}", "\n\n", text).strip()
    if not text:
        raise DocumentParseError("Could not extract any text from the uploaded document.")
    if len(text) > max_chars:
        text = text[:max_chars] + "\n\n[...truncated...]"
    return text


def _decode_text(data: bytes) -> str:
    for enc in ("utf-8", "utf-16", "latin-1"):
        try:
            return data.decode(enc)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="ignore")


def _extract_docx(data: bytes) -> str:
    from docx import Document

    doc = Document(io.BytesIO(data))
    parts: list[str] = []
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text.strip())
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _extract_pdf(data: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise DocumentParseError("PDF support requires pypdf. Run: pip install pypdf") from exc
    reader = PdfReader(io.BytesIO(data))
    pages = []
    for page in reader.pages[:30]:
        pages.append(page.extract_text() or "")
    return "\n".join(pages)


def _extract_xlsx(data: bytes) -> str:
    from openpyxl import load_workbook

    wb = load_workbook(io.BytesIO(data), data_only=True, read_only=True)
    lines: list[str] = []
    for sheet in wb.worksheets[:5]:
        lines.append(f"# Sheet: {sheet.title}")
        for i, row in enumerate(sheet.iter_rows(values_only=True)):
            if i > 80:
                lines.append("...truncated...")
                break
            vals = [str(c) for c in row if c is not None and str(c).strip()]
            if vals:
                lines.append(" | ".join(vals))
    return "\n".join(lines)
