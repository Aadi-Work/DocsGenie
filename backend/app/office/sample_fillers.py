"""Fillers that clone the real PS-08 sample Office files and replace content in-place."""

from __future__ import annotations

import re
import shutil
from copy import copy, deepcopy
from datetime import datetime
from pathlib import Path
from typing import Any

from dataclasses import dataclass, field

from docx import Document
from openpyxl import load_workbook
from openpyxl.cell.cell import ILLEGAL_CHARACTERS_RE
from openpyxl.utils import range_boundaries
from openpyxl.utils.cell import column_index_from_string, coordinate_from_string
from openpyxl.worksheet.worksheet import Worksheet


@dataclass
class TemplateSpec:
    id: str = ""
    filler: str | None = None


@dataclass
class RenderResult:
    output_path: str
    template_id: str = ""
    filled_slots: list = field(default_factory=list)
    coverage: float = 1.0


def _sanitize_xml(value: Any) -> Any:
    if not isinstance(value, str):
        return value
    return ILLEGAL_CHARACTERS_RE.sub("", value)


def _as_text(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, list):
        if value and isinstance(value[0], dict):
            return _sanitize_xml("\n".join(" | ".join(str(v) for v in row.values()) for row in value))
        return _sanitize_xml(", ".join(str(x) for x in value))
    if isinstance(value, datetime):
        return value.strftime("%Y-%m-%d")
    return _sanitize_xml(str(value))


def _copy_row_style(ws: Worksheet, src_row: int, dest_row: int, max_col: int = 7) -> None:
    _clone_row_format(ws, src_row, dest_row, 1, max_col)


def _clone_row_format(ws: Worksheet, src_row: int, dest_row: int, min_col: int = 1, max_col: int = 7) -> None:
    """Clone cell xf, number format, height, and horizontal merges from a template row."""
    if src_row == dest_row:
        return
    for col in range(min_col, max_col + 1):
        src = ws.cell(src_row, col)
        dest = ws.cell(dest_row, col)
        if type(dest).__name__ == "MergedCell":
            continue
        if type(src).__name__ == "MergedCell":
            continue
        if getattr(src, "has_style", False):
            dest._style = copy(src._style)
            dest.number_format = src.number_format
    src_dim = ws.row_dimensions.get(src_row)
    if src_dim is not None and src_dim.height is not None:
        ws.row_dimensions[dest_row].height = src_dim.height
    for rng in list(ws.merged_cells.ranges):
        min_c, min_r, max_c, max_r = range_boundaries(str(rng))
        if min_r == src_row and max_r == src_row and min_c >= min_col and max_c <= max_col:
            already = any(
                range_boundaries(str(existing)) == (min_c, dest_row, max_c, dest_row)
                for existing in ws.merged_cells.ranges
            )
            if already:
                continue
            try:
                ws.merge_cells(
                    start_row=dest_row,
                    start_column=min_c,
                    end_row=dest_row,
                    end_column=max_c,
                )
            except Exception:
                pass


def _safe_set(ws: Worksheet, row: int, col: int, value: Any) -> None:
    """Set a cell value; if target is inside a merge, write the merge anchor instead."""
    value = _sanitize_xml(value)
    cell = ws.cell(row, col)
    if type(cell).__name__ != "MergedCell":
        cell.value = value
        return
    for merged in ws.merged_cells.ranges:
        if merged.min_row <= row <= merged.max_row and merged.min_col <= col <= merged.max_col:
            anchor = ws.cell(merged.min_row, merged.min_col)
            if type(anchor).__name__ != "MergedCell":
                anchor.value = value
            return


def _unmerge_area(
    ws: Worksheet,
    min_row: int,
    max_row: int,
    min_col: int = 1,
    max_col: int = 20,
) -> None:
    """Remove merges that intersect a write area so per-column fills work."""
    victims = [
        str(rng)
        for rng in list(ws.merged_cells.ranges)
        if not (
            rng.max_row < min_row
            or rng.min_row > max_row
            or rng.max_col < min_col
            or rng.min_col > max_col
        )
    ]
    for coord in victims:
        try:
            ws.unmerge_cells(coord)
        except Exception:
            pass


def _safe_clear_row(ws: Worksheet, row: int, max_col: int = 7) -> None:
    for c in range(1, max_col + 1):
        _safe_set(ws, row, c, None)


def _row_blob(ws: Worksheet, row: int, max_col: int = 8) -> str:
    parts = []
    for c in range(1, max_col + 1):
        val = ws.cell(row, c).value
        if val not in (None, ""):
            parts.append(str(val))
    return " ".join(parts)


def _row_has(ws: Worksheet, row: int, *needles: str) -> bool:
    blob = _row_blob(ws, row).lower()
    return all(n.lower() in blob for n in needles)


def _find_row(ws: Worksheet, *needles: str, start: int = 1, end: int = 80) -> int | None:
    cap = min(end, 200)
    for r in range(start, cap + 1):
        if _row_has(ws, r, *needles):
            return r
    return None


def _sheet_named(wb, wanted: str):
    if wanted in wb.sheetnames:
        return wb[wanted]
    low = (wanted or "").strip().lower()
    for name in wb.sheetnames:
        if low and (low in name.lower() or name.lower() in low):
            return wb[name]
    return wb[wb.sheetnames[0]]


def _coord_rc(coord: str) -> tuple[int, int]:
    col_letter, row = coordinate_from_string(coord)
    return row, column_index_from_string(col_letter)


def _normalize_bfl_functions(raw: Any, keys: list[str]) -> list[dict[str, str]]:
    """Accept list[dict], list[str], or a pipe-row blob. Never iterate a string by character."""
    if raw in (None, "", []):
        return []
    if isinstance(raw, dict):
        items: list[Any] = [raw]
    elif isinstance(raw, str):
        text = raw.strip()
        items = [ln.strip() for ln in text.splitlines() if ln.strip()] if "\n" in text else ([text] if text else [])
    elif isinstance(raw, list):
        items = raw
    else:
        items = []

    out: list[dict[str, str]] = []
    for row in items[:80]:
        if isinstance(row, str):
            parts = [p.strip() for p in row.split("|")]
            parsed = {keys[j]: (parts[j] if j < len(parts) else "") for j in range(len(keys))}
            extra = parts[len(keys) :]
            if extra:
                parsed["steps"] = " ".join(x for x in [parsed.get("steps", ""), *extra] if x).strip()
            row = parsed
        if not isinstance(row, dict):
            continue
        cleaned = {k: str(row.get(k, "") or "").strip() for k in keys}
        if not cleaned.get("process"):
            cleaned["process"] = str(
                row.get("name") or row.get("function") or row.get("business_process") or ""
            ).strip()
        if not cleaned.get("description"):
            cleaned["description"] = str(row.get("desc") or row.get("summary") or "").strip()
        if cleaned.get("process") or cleaned.get("description"):
            out.append(cleaned)
    return out


def _sheet_has_mustache(ws: Worksheet) -> bool:
    for r in range(1, min(40, (ws.max_row or 1) + 1) + 5):
        if "{{" in _row_blob(ws, r):
            return True
    return False


def _insert_rows_preserve(ws: Worksheet, idx: int, amount: int) -> None:
    """insert_rows does not shift row_dimensions; keep section/header heights."""
    if amount <= 0:
        return
    last = min(ws.max_row or idx, idx + 400)
    heights = {r: ws.row_dimensions[r].height for r in range(idx, last + 1)}
    ws.insert_rows(idx, amount)
    for old_r, height in sorted(heights.items(), reverse=True):
        ws.row_dimensions[old_r + amount].height = height


def _expand_template_row(ws: Worksheet, template_row: int, count: int, max_col: int = 7) -> None:
    extra = max(count, 1) - 1
    if extra <= 0:
        return
    _insert_rows_preserve(ws, template_row + 1, extra)
    for offset in range(extra):
        _clone_row_format(ws, template_row, template_row + 1 + offset, 1, max_col)


def _as_excel_date(value: Any) -> Any:
    if value in (None, ""):
        return None
    text = str(value).strip()
    try:
        return datetime.strptime(text[:10], "%Y-%m-%d")
    except ValueError:
        return value


def _normalize_actions(actions: Any) -> list[dict[str, str]]:
    if isinstance(actions, str):
        actions = [ln for ln in actions.splitlines() if ln.strip()]
    out: list[dict[str, str]] = []
    for row in actions or []:
        if isinstance(row, str):
            parts = [p.strip() for p in row.split("|")]
            row = {
                "action": parts[0] if parts else row,
                "owner": parts[1] if len(parts) > 1 else "",
                "due_date": parts[2] if len(parts) > 2 else "",
                "status": parts[3] if len(parts) > 3 else "Open",
                "remarks": parts[4] if len(parts) > 4 else "",
            }
        if not isinstance(row, dict):
            continue
        out.append(row)
    return out


def _normalize_summary(summary: Any) -> list[str]:
    if isinstance(summary, str):
        summary = [s.strip() for s in summary.replace(";", "\n").split("\n") if s.strip()]
    items: list[str] = []
    for item in summary or []:
        if isinstance(item, dict):
            item = item.get("text") or item.get("topic") or item.get("details") or " | ".join(
                str(v) for v in item.values() if v not in (None, "")
            )
        text = str(item).strip()
        if text:
            items.append(text)
    return items


def _write_action_row(ws: Worksheet, r: int, i: int, row: dict[str, Any]) -> None:
    _safe_set(ws, r, 1, i + 1)
    _safe_set(ws, r, 2, row.get("action") or row.get("task") or "")
    _safe_set(ws, r, 3, row.get("owner") or row.get("responsibility") or "")
    due = row.get("due_date") or row.get("target_date") or ""
    _safe_set(ws, r, 4, _as_excel_date(due) if due else None)
    closure = row.get("closure_date") or ""
    _safe_set(ws, r, 5, _as_excel_date(closure) if closure else None)
    _safe_set(ws, r, 6, row.get("status") or "Open")
    _safe_set(ws, r, 7, row.get("remarks") or "")


def _apply_table_body_style(ws: Worksheet, header_row: int, dest_row: int, max_col: int = 7) -> None:
    """Keep header chrome; give body rows the same thin grid without the header fill/bold."""
    from openpyxl.styles import Alignment, Border, Side

    thin = Side(style="thin")
    for c in range(1, max_col + 1):
        header = ws.cell(header_row, c)
        dest = ws.cell(dest_row, c)
        if type(dest).__name__ == "MergedCell":
            continue
        src_border = header.border
        if src_border and any(
            getattr(getattr(src_border, side, None), "style", None)
            for side in ("left", "right", "top", "bottom")
        ):
            dest.border = copy(src_border)
        else:
            dest.border = Border(left=thin, right=thin, top=thin, bottom=thin)
        dest.alignment = Alignment(
            wrap_text=True,
            vertical="center",
            horizontal=header.alignment.horizontal if header.alignment else None,
        )
    if ws.row_dimensions[dest_row].height is None:
        ws.row_dimensions[dest_row].height = max(ws.row_dimensions[header_row].height or 18, 18)


def fill_mom_sample(
    spec: TemplateSpec,
    template_path: Path,
    context: dict[str, Any],
    output_path: Path,
) -> RenderResult:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(template_path, output_path)
    wb = load_workbook(output_path)
    ws = wb[wb.sheetnames[0]]

    for coord in ("B5", "E5", "B6", "E6", "B7", "E7"):
        cell = ws[coord]
        _safe_set(ws, cell.row, cell.column, None)

    date_val = context.get("meeting_date")
    if date_val:
        _safe_set(ws, 5, 2, _as_excel_date(date_val))
    if context.get("purpose"):
        _safe_set(ws, 5, 5, _as_text(context["purpose"]))
    if context.get("prepared_by"):
        _safe_set(ws, 6, 2, _as_text(context["prepared_by"]))
    if context.get("venue"):
        _safe_set(ws, 6, 5, _as_text(context["venue"]))
    if context.get("attendees_ymsli") is not None:
        _safe_set(ws, 7, 2, _as_text(context["attendees_ymsli"]))
    if context.get("attendees_ymesg") is not None:
        _safe_set(ws, 7, 5, _as_text(context["attendees_ymesg"]))

    summary = _normalize_summary(context.get("summary_items"))
    actions = _normalize_actions(context.get("action_items"))
    mustache = _sheet_has_mustache(ws)

    if mustache or _find_row(ws, "action plan after current meeting"):
        # MOM_Template.xlsx: keep header row 14 and clone styled data row 15.
        summary_tpl = _find_row(ws, "{{#discussions}}") or _find_row(ws, "{{details}}")
        if summary_tpl is None:
            sl = _find_row(ws, "meeting summary:")
            if sl:
                nxt = sl + 1
                if "sl.no" in _row_blob(ws, nxt).lower():
                    summary_tpl = nxt + 1
                else:
                    summary_tpl = nxt
        if summary_tpl:
            _expand_template_row(ws, summary_tpl, max(len(summary), 1), max_col=7)
            if summary:
                for i, text in enumerate(summary):
                    r = summary_tpl + i
                    _safe_set(ws, r, 1, i + 1)
                    _safe_set(ws, r, 2, text)
            else:
                _safe_set(ws, summary_tpl, 1, None)
                _safe_set(ws, summary_tpl, 2, None)

        header = _find_row(ws, "action", "responsibility", "target") or _find_row(
            ws, "action", "status", "remarks"
        )
        action_tpl = (header + 1) if header else (_find_row(ws, "{{#action_items}}") or _find_row(ws, "{{task}}"))
        if action_tpl:
            header_row = action_tpl - 1
            _expand_template_row(ws, action_tpl, max(len(actions), 1), max_col=7)
            if actions:
                for i, row in enumerate(actions):
                    r = action_tpl + i
                    _write_action_row(ws, r, i, row)
                    _apply_table_body_style(ws, header_row, r, max_col=7)
            else:
                for c in range(1, 8):
                    _safe_set(ws, action_tpl, c, None)
    else:
        summary_start, summary_end = 11, 24
        style_summary = 11
        for r in range(summary_start, summary_end + 1):
            _safe_clear_row(ws, r, max_col=7)
        for i, text in enumerate(summary[: summary_end - summary_start + 1]):
            r = summary_start + i
            _safe_set(ws, r, 1, i + 1)
            _safe_set(ws, r, 2, str(text))
            _clone_row_format(ws, style_summary, r, 1, 7)

        action_start = 28
        style_action = 28
        if len(actions) > 1:
            _expand_template_row(ws, action_start, len(actions), max_col=7)
        clear_to = max(ws.max_row or action_start, action_start + max(len(actions), 1) - 1)
        for r in range(action_start, clear_to + 1):
            _safe_clear_row(ws, r, max_col=7)
        for i, row in enumerate(actions):
            r = action_start + i
            _write_action_row(ws, r, i, row)
            _clone_row_format(ws, style_action, r, 1, 7)

    wb.save(output_path)
    return RenderResult(
        output_path=str(output_path),
        template_id=spec.id,
        filled_slots=[k for k, v in context.items() if v not in (None, "", [])],
        coverage=1.0,
    )


def fill_poc_sample(
    spec: TemplateSpec,
    template_path: Path,
    context: dict[str, Any],
    output_path: Path,
) -> RenderResult:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(template_path, output_path)
    wb = load_workbook(output_path)

    cover = wb["Cover"]
    if context.get("project_name"):
        cover["B15"] = _as_text(context["project_name"])
    if context.get("workstream"):
        cover["B16"] = _as_text(context["workstream"])

    if "History." in wb.sheetnames and context.get("updated_by"):
        hist = wb["History."]
        hist["D3"] = _as_text(context["updated_by"])
        if context.get("meeting_date") or context.get("document_date"):
            raw = context.get("document_date") or context.get("meeting_date")
            try:
                hist["C3"] = datetime.strptime(str(raw)[:10], "%Y-%m-%d")
            except Exception:
                pass

    sheet_name = "YMVN POC list" if "YMVN POC list" in wb.sheetnames else wb.sheetnames[-2]
    ws = wb[sheet_name]

    # Capture style from first data row before clearing
    style_src_row = 2
    pocs = context.get("pocs") or []
    keys = ["id", "cycle", "subprocess", "title", "description", "prereq", "steps", "decision", "azure_ids"]

    max_clear = max(ws.max_row or 2, 2)
    for r in range(2, min(max_clear, 500) + 1):
        for c in range(1, 10):
            _safe_set(ws, r, c, None)

    for i, row in enumerate(pocs):
        r = 2 + i
        if isinstance(row, str):
            parts = [p.strip() for p in row.split("|")]
            row = {keys[j]: (parts[j] if j < len(parts) else "") for j in range(len(keys))}
        for j, key in enumerate(keys, start=1):
            ws.cell(r, j).value = row.get(key, "")
        _copy_row_style(ws, style_src_row, r, max_col=9)

    wb.save(output_path)
    return RenderResult(
        output_path=str(output_path),
        template_id=spec.id,
        filled_slots=[k for k, v in context.items() if v not in (None, "", [])],
        coverage=1.0,
    )


def _ensure_row_merge(ws: Worksheet, row: int, min_col: int, max_col: int) -> None:
    already = any(
        range_boundaries(str(rng))[:3] == (min_col, row, max_col) and range_boundaries(str(rng))[3] == row
        for rng in ws.merged_cells.ranges
    )
    if already:
        return
    try:
        ws.merge_cells(start_row=row, start_column=min_col, end_row=row, end_column=max_col)
    except Exception:
        pass


def _format_bfl_steps(steps: str) -> str:
    text = _as_text(steps)
    text = re.sub(r"\s*;\s*", "\n", text)
    text = re.sub(r"(?<=\S)\s+(?=\d+[.)]\s+)", "\n", text)
    return text.strip()


def fill_bfl_sample(
    spec: TemplateSpec,
    template_path: Path,
    context: dict[str, Any],
    output_path: Path,
) -> RenderResult:
    """
    Fill real BFL Sample.xlsx:
      Cover C16/C17/C18 + Business Function List
      Keep Level 1/2 C:O merges; one styled Level 3 row per function.
    """
    from openpyxl.styles import Alignment

    from app.office.agent.bfl_intelligence import polish_function_row
    from app.office.layouts import BFL_LAYOUT

    output_path.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(template_path, output_path)
    wb = load_workbook(output_path)

    cover = _sheet_named(wb, BFL_LAYOUT["cover_sheet"])
    for key, coord in BFL_LAYOUT["cover"].items():
        val = context.get(key)
        if val in (None, ""):
            continue
        row, col = _coord_rc(coord)
        _safe_set(cover, row, col, _as_text(val))

    ws = _sheet_named(wb, BFL_LAYOUT["list_sheet"])
    cols = BFL_LAYOUT["columns"]
    keys = list(BFL_LAYOUT["pipe_keys"])
    process = _as_text(context.get("business_process") or "Business Process")
    workstream = _as_text(context.get("workstream") or "Functions")

    norm_rows = _normalize_bfl_functions(context.get("functions"), keys)
    norm_rows = [polish_function_row(row, i + 1, process) for i, row in enumerate(norm_rows)]

    start = int(BFL_LAYOUT["functions_start_row"])
    style_src = start
    sample_end = start
    for r in range(start, 81):
        if ws.cell(r, cols["process"]).value not in (None, "") or ws.cell(r, cols["level"]).value not in (None, ""):
            sample_end = r
    clear_to = max(sample_end, start + max(len(norm_rows), 1) + 2, 12)
    clear_to = min(clear_to, 80)

    # Keep Level 1/2 C:O banners. Only flatten the function body (vertical step merges).
    _unmerge_area(ws, min_row=start, max_row=clear_to, min_col=1, max_col=16)
    for r in range(start, clear_to + 1):
        for c in range(1, 16):
            _safe_set(ws, r, c, None)

    # Level 1 — top process (C3:O3 stays merged)
    _safe_set(ws, 3, cols["id"], 100)
    _safe_set(ws, 3, cols["level"], 1)
    _safe_set(ws, 3, cols["process"], process)
    _ensure_row_merge(ws, 3, cols["process"], 15)

    # Level 2 — mid group (C4:O4 stays merged)
    code = "55"
    m = re.match(r"(\d+)", process)
    if m:
        code = m.group(1)
    _safe_set(ws, 4, cols["id"], 110)
    _safe_set(ws, 4, cols["level"], 2)
    _safe_set(ws, 4, cols["process"], f"{code}.10 {workstream}")
    _ensure_row_merge(ws, 4, cols["process"], 15)

    wrap = Alignment(wrap_text=True, vertical="top")
    for i, row in enumerate(norm_rows):
        r = start + i
        steps = _format_bfl_steps(row.get("steps", ""))
        if r != style_src:
            _clone_row_format(ws, style_src, r, 1, 15)
        _apply_table_body_style(ws, 2, r, 15)

        _safe_set(ws, r, cols["id"], 111 + i)
        _safe_set(ws, r, cols["level"], 3)
        _safe_set(ws, r, cols["process"], row.get("process", ""))
        _safe_set(ws, r, cols["description"], row.get("description", ""))
        _safe_set(ws, r, cols["steps"], steps)
        if cols.get("item_type"):
            _safe_set(ws, r, cols["item_type"], "Function")
        _safe_set(ws, r, cols["input"], row.get("input", ""))
        _safe_set(ws, r, cols["output"], row.get("output", ""))
        _safe_set(ws, r, cols["department"], row.get("department") or "Finance")
        _safe_set(ws, r, cols["frequency"], row.get("frequency") or "Ondemand")
        _safe_set(ws, r, cols["manual_auto"], row.get("manual_auto") or "Manual")
        _safe_set(ws, r, cols["type"], row.get("type") or "Screen")
        _safe_set(ws, r, cols["module"], row.get("module") or "Fixed Assets")
        _safe_set(ws, r, cols["fit_gap"], row.get("fit_gap") or "Fit")

        for col in (cols["process"], cols["description"], cols["steps"]):
            cell = ws.cell(r, col)
            if type(cell).__name__ != "MergedCell":
                cell.alignment = wrap
        nlines = max(1, steps.count("\n") + 1)
        ws.row_dimensions[r].height = min(20 + nlines * 13, 96)

    wb.save(output_path)
    return RenderResult(
        output_path=str(output_path),
        template_id=spec.id,
        filled_slots=[k for k, v in context.items() if v not in (None, "", [])],
        coverage=1.0 if norm_rows else 0.4,
    )


def _replace_paragraph_text(paragraph, new_text: str) -> None:
    if not paragraph.runs:
        paragraph.add_run(new_text)
        return
    paragraph.runs[0].text = new_text
    for run in paragraph.runs[1:]:
        run.text = ""


def _set_oxml_paragraph_text(p_el, text: str, qn) -> None:
    """Write text into an existing paragraph element, keeping pPr / rPr."""
    from docx.oxml import OxmlElement

    value = "" if text is None else str(text)
    runs = [child for child in p_el.iterchildren() if child.tag == qn("w:r")]
    if not runs:
        run = OxmlElement("w:r")
        node = OxmlElement("w:t")
        node.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        node.text = value
        run.append(node)
        p_el.append(run)
        return
    first = True
    for run in runs:
        texts = [child for child in run.iterchildren() if child.tag == qn("w:t")]
        if first:
            if texts:
                texts[0].text = value
                texts[0].set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
                for extra in texts[1:]:
                    extra.text = ""
            else:
                node = OxmlElement("w:t")
                node.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
                node.text = value
                run.append(node)
            first = False
        else:
            for node in texts:
                node.text = ""


def _set_cell_text(cell, text: str) -> None:
    """Keep cell borders/shading/font; only replace the visible text."""
    value = "" if text is None else str(text)
    if not cell.paragraphs:
        cell.add_paragraph(value)
        return
    _replace_paragraph_text(cell.paragraphs[0], value)
    for extra in list(cell.paragraphs[1:]):
        parent = extra._element.getparent()
        if parent is not None:
            parent.remove(extra._element)


def _clone_table_row(table, src_idx: int):
    new_tr = deepcopy(table.rows[src_idx]._tr)
    table._tbl.append(new_tr)
    row = table.rows[-1]
    for cell in row.cells:
        _set_cell_text(cell, "")
    return row


def _fill_styled_table(table, rows: list[list[str]], style_row_idx: int = 1) -> None:
    """Replace data rows while cloning the sample data-row formatting."""
    if not rows:
        while len(table.rows) > 1:
            table._tbl.remove(table.rows[-1]._tr)
        return
    src = style_row_idx if len(table.rows) > style_row_idx else max(0, len(table.rows) - 1)
    while len(table.rows) < len(rows) + 1:
        _clone_table_row(table, src)
    for i, values in enumerate(rows):
        cells = table.rows[i + 1].cells
        for col, value in enumerate(values):
            if col < len(cells):
                _set_cell_text(cells[col], value)
    while len(table.rows) > len(rows) + 1:
        table._tbl.remove(table.rows[-1]._tr)


def _para_plain(p_el, qn) -> str:
    return "".join((node.text or "") for node in p_el.iter() if node.tag == qn("w:t")).strip()


def _brd_para_role(text: str) -> str:
    raw = (text or "").strip()
    if not raw:
        return "blank"
    low = raw.lower()
    if low.startswith("area path"):
        return "area"
    if low.rstrip(":") == "description":
        return "desc_label"
    if low.startswith("acceptance"):
        return "acc_label"
    if low.startswith("process flow") or low.rstrip(":") in {"flow", "process steps", "steps"}:
        return "flow_label"
    if "type:" in low and ("|" in raw or "board state" in low):
        return "title"
    return "body"


def _is_page_break(p_el, qn) -> bool:
    for br in p_el.iter(qn("w:br")):
        if br.get(qn("w:type")) == "page":
            return True
    return False


def _iter_item_blocks(body, qn) -> list[list[Any]]:
    """Each sample item after the summary table, including page-break paragraphs."""
    blocks: list[list[Any]] = []
    current: list[Any] = []
    tables_seen = 0
    started = False
    for child in list(body.iterchildren()):
        if child.tag == qn("w:tbl"):
            tables_seen += 1
            continue
        if child.tag == qn("w:sectPr"):
            break
        if tables_seen < 2:
            continue
        text = _para_plain(child, qn) if child.tag == qn("w:p") else ""
        if child.tag == qn("w:p") and _brd_para_role(text) == "title":
            if current:
                blocks.append(current)
            current = [child]
            started = True
        elif started:
            current.append(child)
    if current:
        blocks.append(current)
    return blocks


def _number_flow_steps(lines: list[str]) -> list[str]:
    out: list[str] = []
    for i, line in enumerate(lines, start=1):
        text = str(line).strip()
        if not text:
            continue
        if re.match(r"^\d+[.)]\s+", text):
            out.append(text)
        else:
            out.append(f"{i}. {text}")
    return out


def _fill_item_block(block: list[Any], item: dict[str, Any], area_path: str, overview: str, idx: int, qn) -> None:
    """Overwrite sample text in a cloned item block; keep page breaks and paragraph XML."""
    desc = item.get("description") or ""
    if idx == 0 and overview and item.get("type") == "Epic" and not desc:
        desc = overview
    elif idx == 0 and overview and item.get("type") == "Epic" and overview not in desc:
        desc = f"{overview}\n\n{desc}".strip() if desc else overview
    desc_lines = _split_brd_lines(desc or "(Description to be detailed.)")
    acc_lines = _split_brd_lines(str(item.get("acceptance") or ""))
    flow_lines = _number_flow_steps(
        [str(x).strip() for x in (item.get("flow") or []) if str(x).strip()]
        if isinstance(item.get("flow"), list)
        else _split_brd_lines(str(item.get("flow") or ""))
    )
    title = (
        f"{item['name']} | Type: {item['type']} | "
        f"Board state: New | Applicable company:"
    )

    mode = "pre"
    desc_i = 0
    acc_i = 0
    flow_i = 0
    saw_flow = False
    last_by_mode: dict[str, Any] = {"desc": None, "acc": None, "flow": None}
    label_el: dict[str, Any] = {"desc": None, "acc": None, "flow": None}
    body_tmpl = None
    for el in block:
        if el.tag != qn("w:p"):
            continue
        if _is_page_break(el, qn) and not _para_plain(el, qn):
            continue
        role = _brd_para_role(_para_plain(el, qn))
        if role == "title":
            _set_oxml_paragraph_text(el, title, qn)
            mode = "pre"
        elif role == "area":
            _set_oxml_paragraph_text(el, f"Area Path: {area_path}", qn)
        elif role == "desc_label":
            _set_oxml_paragraph_text(el, "Description:", qn)
            mode = "desc"
            label_el["desc"] = el
        elif role == "acc_label":
            _set_oxml_paragraph_text(el, "Acceptance Criteria:", qn)
            mode = "acc"
            label_el["acc"] = el
        elif role == "flow_label":
            _set_oxml_paragraph_text(el, "Process Flow:", qn)
            mode = "flow"
            saw_flow = True
            label_el["flow"] = el
        elif role == "blank":
            continue
        elif mode == "desc":
            _set_oxml_paragraph_text(el, desc_lines[desc_i] if desc_i < len(desc_lines) else "", qn)
            last_by_mode["desc"] = el
            body_tmpl = body_tmpl or el
            desc_i += 1
        elif mode == "acc":
            _set_oxml_paragraph_text(el, acc_lines[acc_i] if acc_i < len(acc_lines) else "", qn)
            last_by_mode["acc"] = el
            body_tmpl = body_tmpl or el
            acc_i += 1
        elif mode == "flow":
            _set_oxml_paragraph_text(el, flow_lines[flow_i] if flow_i < len(flow_lines) else "", qn)
            last_by_mode["flow"] = el
            body_tmpl = body_tmpl or el
            flow_i += 1

    extra: list[Any] = []

    def _insert_after(anchor: Any, text: str) -> Any:
        src = body_tmpl if body_tmpl is not None else anchor
        new_p = deepcopy(src)
        _set_oxml_paragraph_text(new_p, text, qn)
        anchor.addnext(new_p)
        extra.append(new_p)
        return new_p

    def _more(mode_key: str, lines: list[str], used: int) -> None:
        anchor = last_by_mode[mode_key] or label_el[mode_key]
        if anchor is None:
            return
        for line in lines[used:]:
            anchor = _insert_after(anchor, line)
            last_by_mode[mode_key] = anchor

    _more("desc", desc_lines, desc_i)
    _more("acc", acc_lines, acc_i)
    if flow_lines:
        if not saw_flow:
            anchor = last_by_mode["acc"] or last_by_mode["desc"] or label_el["acc"] or label_el["desc"] or block[-1]
            label_p = _insert_after(anchor, "Process Flow:")
            label_el["flow"] = label_p
            last_by_mode["flow"] = label_p
            used = 0
        else:
            used = flow_i
        _more("flow", flow_lines, used)
    block.extend(extra)


def _insert_cloned_block(template_block: list[Any], sect, qn) -> list[Any]:
    cloned = [deepcopy(el) for el in template_block]
    for el in cloned:
        if sect is not None:
            sect.addprevious(el)
        else:
            template_block[0].getparent().append(el)
    return cloned


def _snapshot_brd_templates(body, qn) -> dict[str, Any]:
    """Keep one styled paragraph per role from the sample item block after table 2."""
    templates: dict[str, Any] = {}
    tables_seen = 0
    for child in list(body.iterchildren()):
        if child.tag == qn("w:tbl"):
            tables_seen += 1
            continue
        if tables_seen < 2 or child.tag != qn("w:p"):
            continue
        role = _brd_para_role(_para_plain(child, qn))
        templates.setdefault(role, child)
        if {"title", "area", "desc_label", "body", "acc_label"} <= set(templates):
            break
    return templates


def _append_styled_paragraph(doc, template_el, text: str, qn, *, bold: bool | None = None) -> None:
    from docx.oxml import OxmlElement
    from docx.shared import Pt

    body = doc.element.body
    sect = body.find(qn("w:sectPr"))
    value = "" if text is None else str(text)
    if template_el is not None:
        new_p = deepcopy(template_el)
        _set_oxml_paragraph_text(new_p, value, qn)
        if sect is not None:
            sect.addprevious(new_p)
        else:
            body.append(new_p)
        return
    para = doc.add_paragraph()
    run = para.add_run(value)
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    if bold:
        run.bold = True
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), "Calibri")
    rfonts.set(qn("w:hAnsi"), "Calibri")


def _split_brd_lines(text: str) -> list[str]:
    raw = (text or "").strip()
    if not raw:
        return [""]
    if " ; " in raw:
        parts = [p.strip() for p in raw.split(" ; ") if p.strip()]
        return parts or [raw]
    if "\n" in raw:
        return [ln.strip() for ln in raw.splitlines() if ln.strip()] or [raw]
    numbered = re.split(r"(?=\s*\d+[.)]\s)", raw)
    parts = [p.strip() for p in numbered if p.strip()]
    if len(parts) > 1:
        return parts
    return [raw]


def fill_brd_sample(
    spec: TemplateSpec,
    template_path: Path,
    context: dict[str, Any],
    output_path: Path,
) -> RenderResult:
    """
    High-accuracy fill for BRD Sample.docx:
      - Header (title, process, prepared by)
      - Revision table (formatting preserved)
      - Items summary table (cloned data-row formatting)
      - Rebuild detail sections using the sample paragraph styles
        (keeps w:sectPr so page setup survives the rewrite)
    """
    from docx.oxml.ns import qn

    from app.office.agent.brd_intelligence import generate_item_flow, polish_brd_context

    output_path.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(template_path, output_path)
    doc = Document(str(output_path))
    context = polish_brd_context(context)

    doc_code = _as_text(context.get("doc_code") or "SE52: Business Requirement Document")
    process_name = _as_text(context.get("process_name") or "")
    prepared_by = _as_text(context.get("prepared_by") or "YMSLI")
    overview = _as_text(context.get("overview") or "")
    area_path = _as_text(context.get("area_path") or "YNS-FnO-ERP")
    raw_date = _as_text(context.get("document_date") or "")
    rev_date = raw_date
    # Prefer DD/MM/YYYY to match sample revision style
    try:
        rev_date = datetime.strptime(raw_date[:10], "%Y-%m-%d").strftime("%d/%m/%Y")
    except Exception:
        pass

    items_in = context.get("items") or []
    items: list[dict[str, str]] = []
    for item in items_in:
        if isinstance(item, str):
            parts = [p.strip() for p in item.split("|")]
            item = {
                "type": parts[0] if parts else "Feature",
                "name": parts[1] if len(parts) > 1 else item,
                "description": parts[2] if len(parts) > 2 else "",
                "acceptance": parts[3] if len(parts) > 3 else "",
                "flow": parts[4] if len(parts) > 4 else "",
            }
        if not isinstance(item, dict):
            continue
        typ = str(item.get("type") or "Feature").strip()
        name = str(item.get("name") or "").strip()
        desc = str(item.get("description") or item.get("desc") or "").strip()
        acc = str(item.get("acceptance") or item.get("acceptance_criteria") or "").strip()
        if not name:
            continue
        # Normalize type casing
        low = typ.lower().replace("_", " ")
        if "epic" in low:
            typ = "Epic"
        elif "user" in low and "story" in low:
            typ = "User Story"
        elif "feature" in low:
            typ = "Feature"
        if not name.startswith("[") and typ in {"Epic", "Feature"}:
            # Match sample naming: Epic uses process title; Features get [YNS] prefix only
            if typ == "Epic" and process_name:
                name = f"[YNS] {process_name}"
            else:
                name = f"[YNS] {name}"
        flow = item.get("flow") or item.get("process_flow") or []
        if isinstance(flow, str):
            flow = [ln.strip() for ln in flow.splitlines() if ln.strip()] if flow.strip() else []
        items.append(
            {"type": typ, "name": name, "description": desc, "acceptance": acc, "flow": flow}
        )
    for row in items:
        if not row.get("flow"):
            row["flow"] = generate_item_flow(row, items, overview)

    # Header paragraphs
    seen_title = False
    for para in doc.paragraphs:
        t = para.text.strip()
        if t.startswith("SE52:") or t == "SE52: Business Requirement Document":
            _replace_paragraph_text(para, doc_code)
            seen_title = True
        elif seen_title and process_name and (
            "Capital Asset Acquisition" in t
            or (len(t) < 100 and re.match(r"^\d", t))
            or t.lower().startswith("process")
        ):
            _replace_paragraph_text(para, process_name)
            seen_title = False
        elif t.startswith("Prepared by:"):
            _replace_paragraph_text(para, f"Prepared by: {prepared_by}")

    # Revision table — keep sample cell formatting
    if doc.tables:
        rev = doc.tables[0]
        if len(rev.rows) > 1:
            cells = rev.rows[1].cells
            _set_cell_text(cells[0], "1.0")
            if len(cells) > 1:
                _set_cell_text(cells[1], rev_date)
            if len(cells) > 2:
                _set_cell_text(cells[2], prepared_by)
            if len(cells) > 4:
                _set_cell_text(cells[4], "Updated via TemplateHub Agent")

    # Summary table — clone the sample data row instead of add_row() (which copies the header)
    if len(doc.tables) > 1:
        summary = doc.tables[1]
        _fill_styled_table(
            summary,
            [[str(i), item["type"], item["name"], area_path] for i, item in enumerate(items, start=1)],
        )

    # Items count heading
    for para in doc.paragraphs:
        if para.text.strip().startswith("Items (") and "Summary" in para.text:
            _replace_paragraph_text(para, f"Items ({len(items)}) Summary:")
            break

    # Fill detail pages in place: clone the sample item block (page breaks included).
    body = doc.element.body
    sect = body.find(qn("w:sectPr"))
    blocks = _iter_item_blocks(body, qn)
    if blocks:
        template_block = [deepcopy(el) for el in blocks[0]]
        for extra in blocks[1:]:
            for el in extra:
                parent = el.getparent()
                if parent is not None:
                    parent.remove(el)
        for idx, item in enumerate(items):
            if idx == 0:
                target = blocks[0]
            else:
                target = _insert_cloned_block(template_block, sect, qn)
            _fill_item_block(target, item, area_path, overview, idx, qn)
    else:
        templates = _snapshot_brd_templates(body, qn)
        title_tmpl = templates.get("title")
        area_tmpl = templates.get("area")
        desc_label_tmpl = templates.get("desc_label")
        body_tmpl = templates.get("body")
        acc_label_tmpl = templates.get("acc_label")
        blank_tmpl = templates.get("blank")
        flow_label_tmpl = templates.get("flow_label")
        for idx, item in enumerate(items):
            _append_styled_paragraph(doc, blank_tmpl, "", qn)
            title = (
                f"{item['name']} | Type: {item['type']} | "
                f"Board state: New | Applicable company:"
            )
            _append_styled_paragraph(doc, title_tmpl, title, qn, bold=True)
            _append_styled_paragraph(doc, area_tmpl, f"Area Path: {area_path}", qn)
            _append_styled_paragraph(doc, desc_label_tmpl, "Description:", qn, bold=True)
            desc = item["description"]
            if idx == 0 and overview and item["type"] == "Epic" and not desc:
                desc = overview
            elif idx == 0 and overview and item["type"] == "Epic" and overview not in desc:
                desc = f"{overview}\n\n{desc}".strip() if desc else overview
            for line in _split_brd_lines(desc or "(Description to be detailed.)"):
                _append_styled_paragraph(doc, body_tmpl, line, qn)
            if item.get("acceptance"):
                _append_styled_paragraph(doc, acc_label_tmpl, "Acceptance Criteria:", qn, bold=True)
                for line in _split_brd_lines(str(item["acceptance"])):
                    _append_styled_paragraph(doc, body_tmpl, line, qn)
            flow_lines = [str(x).strip() for x in (item.get("flow") or []) if str(x).strip()]
            if flow_lines:
                _append_styled_paragraph(doc, flow_label_tmpl or acc_label_tmpl, "Process Flow:", qn, bold=True)
                for line in flow_lines:
                    _append_styled_paragraph(doc, body_tmpl, line, qn)

    doc.save(str(output_path))
    return RenderResult(
        output_path=str(output_path),
        template_id=spec.id,
        filled_slots=[k for k, v in context.items() if v not in (None, "", [])],
        coverage=1.0 if items else 0.4,
    )


def fill_sample_ppt(
    spec: TemplateSpec,
    template_path: Path,
    context: dict[str, Any],
    output_path: Path,
) -> RenderResult:
    from app.office.ppt_pipeline import fill_sample_ppt_file

    fill_sample_ppt_file(Path(template_path), Path(output_path), context)
    return RenderResult(
        output_path=str(output_path),
        template_id=spec.id,
        filled_slots=[k for k, v in context.items() if v not in (None, "", [])],
        coverage=1.0,
    )


FILLERS = {
    "mom_sample": fill_mom_sample,
    "poc_sample": fill_poc_sample,
    "bfl_sample": fill_bfl_sample,
    "brd_sample": fill_brd_sample,
    "sample_ppt": fill_sample_ppt,
}


def render_with_filler(
    spec: TemplateSpec,
    template_path: Path,
    context: dict[str, Any],
    output_path: Path,
) -> RenderResult | None:
    if not spec.filler:
        return None
    fn = FILLERS.get(spec.filler)
    if not fn:
        raise KeyError(f"Unknown filler: {spec.filler}")
    return fn(spec, template_path, context, output_path)
