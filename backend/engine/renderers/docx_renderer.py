"""
DOCX renderer.

Text is written *into existing runs* wherever possible, because a run carries
the font, size, colour and language of the template. New runs are cloned from
a neighbouring run's rPr rather than styled from scratch, and new table rows
are deep-copied from the template row.
"""

from __future__ import annotations

import copy as _copy
import re
from typing import Any, Dict, List, Optional

from docx import Document
from docx.table import _Row, Table
from docx.text.paragraph import Paragraph

from ..ir import DocType
from ..logging_config import get_logger
from ..mapper import FillInstruction, FillPlan
from ..normalize import to_display
from .base import BaseRenderer, RenderResult, WriteRecord, clear_if_unresolved_placeholder, register

log = get_logger("renderers.docx")

PLACEHOLDER_RE = re.compile(r"\{\{\s*([\w.\-]+)\s*\}\}|\$\{\s*([\w.\-]+)\s*\}|<<\s*([\w.\- ]+)\s*>>")


class DocxRenderer(BaseRenderer):
    doc_type = DocType.DOCX

    def render(self, template_path: str, output_path: str, plan: FillPlan,
              clear_unresolved: bool = True) -> RenderResult:
        out = self.prepare_output(template_path, output_path)
        doc = Document(out)
        result = RenderResult(output_path=out)

        for ins in plan.writable():
            if ins.kind == "field":
                result.records.append(self._write_field(doc, ins))
            else:
                rec, added = self._write_table(doc, ins)
                result.records.append(rec)
                result.rows_added += added

        if clear_unresolved:
            cleared = self._clear_unresolved_placeholders(doc)
            if cleared:
                log.info("Cleared %d unresolved placeholder(s) that had no fillable value", cleared)

        doc.save(out)
        return result

    def _clear_unresolved_placeholders(self, doc: Document) -> int:
        n = 0
        for para in doc.paragraphs:
            n += self._clear_paragraph(para)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        n += self._clear_paragraph(para)
        return n

    def _clear_paragraph(self, para: Paragraph) -> int:
        full = "".join(r.text for r in para.runs) or para.text
        cleared = clear_if_unresolved_placeholder(full)
        if cleared == full:
            return 0
        self._set_paragraph_text(para, cleared or "")
        return 1

    # ------------------------------------------------------------------
    def _target_repr(self, t: Dict[str, Any]) -> str:
        if t.get("kind") == "table_cell":
            return f"table[{t.get('table_index')}]!r{t.get('row')}c{t.get('col')}"
        if "para_index" in t:
            return f"paragraph[{t.get('para_index')}]"
        return str(t)

    def _write_field(self, doc: Document, ins: FillInstruction) -> WriteRecord:
        t = ins.target
        rec = WriteRecord(ins.node_id, ins.role, self._target_repr(t), ins.value)
        text = to_display(ins.value)
        try:
            if t.get("kind") == "table_cell":
                table = doc.tables[int(t["table_index"])]
                cell = table.rows[int(t["row"])].cells[int(t["col"])]
                self._set_cell_text(cell, text)
                rec.rows_written = 1
                return rec

            para = doc.paragraphs[int(t["para_index"])]
            mode = t.get("mode", "replace")
            if mode == "after_colon":
                self._append_run(para, (" " if not para.text.endswith(" ") else "") + text)
            elif mode == "placeholder":
                self._replace_placeholder(para, text)
            else:
                self._set_paragraph_text(para, text)
            rec.rows_written = 1
        except Exception as exc:
            rec.status = "failed"
            rec.message = f"{type(exc).__name__}: {exc}"
        return rec

    # ------------------------------------------------------------------
    def _write_table(self, doc: Document, ins: FillInstruction):
        t = ins.target
        rec = WriteRecord(ins.node_id, ins.role, f"table[{t.get('table_index')}]")
        added = 0
        try:
            table = doc.tables[int(t["table_index"])]
            rows = ins.rows or []
            col_map = {f: int(h.get("col")) for f, h in (t.get("columns") or {}).items()
                       if h.get("col") is not None}
            template_row_idx = int(t.get("template_row", 1))

            blank = self._blank_rows(table, template_row_idx)
            for i, row_data in enumerate(rows):
                idx = template_row_idx + i
                if idx >= len(table.rows) or (i >= blank and blank >= 0):
                    if idx >= len(table.rows):
                        self._clone_row(table, template_row_idx)
                        added += 1
                row = table.rows[min(idx, len(table.rows) - 1)]
                for fname, value in row_data.items():
                    ci = col_map.get(fname)
                    if ci is None or ci >= len(row.cells):
                        continue
                    self._set_cell_text(row.cells[ci], to_display(value))
            rec.rows_written = len(rows)
        except Exception as exc:
            rec.status = "failed"
            rec.message = f"{type(exc).__name__}: {exc}"
        return rec, added

    # ------------------------------------------------------------------
    # low-level, formatting-preserving writes
    # ------------------------------------------------------------------
    def _set_paragraph_text(self, para: Paragraph, text: str) -> None:
        if para.runs:
            para.runs[0].text = text
            for r in para.runs[1:]:
                r.text = ""
        else:
            para.add_run(text)

    def _append_run(self, para: Paragraph, text: str) -> None:
        if para.runs:
            src = para.runs[-1]
            new = para.add_run(text)
            if src._element.rPr is not None:
                new._element.insert(0, _copy.deepcopy(src._element.rPr))
        else:
            para.add_run(text)

    def _replace_placeholder(self, para: Paragraph, text: str) -> None:
        """Placeholders often straddle runs; rebuild from the joined text once."""
        full = "".join(r.text for r in para.runs) or para.text
        new_text = PLACEHOLDER_RE.sub(text, full, count=1)
        self._set_paragraph_text(para, new_text)

    def _set_cell_text(self, cell, text: str) -> None:
        lines = str(text).split("\n")
        para = cell.paragraphs[0]
        self._set_paragraph_text(para, lines[0])
        for extra in lines[1:]:
            new_p = _copy.deepcopy(para._p)
            para._p.addnext(new_p)
            para = Paragraph(new_p, cell)
            self._set_paragraph_text(para, extra)

    def _blank_rows(self, table: Table, start_idx: int) -> int:
        n = 0
        for row in list(table.rows)[start_idx:]:
            if any(c.text.strip() for c in row.cells):
                break
            n += 1
        return n

    def _clone_row(self, table: Table, template_row_idx: int) -> _Row:
        """Deep-copy of the template row: borders, shading, widths and all."""
        tmpl = table.rows[template_row_idx]._tr
        new_tr = _copy.deepcopy(tmpl)
        tmpl.addnext(new_tr) if template_row_idx == len(table.rows) - 1 else \
            table.rows[-1]._tr.addnext(new_tr)
        row = table.rows[-1]
        for cell in row.cells:
            for para in cell.paragraphs:
                if para.runs:
                    para.runs[0].text = ""
                    for r in para.runs[1:]:
                        r.text = ""
        return row


register("docx", DocxRenderer)
register("dotx", DocxRenderer)
