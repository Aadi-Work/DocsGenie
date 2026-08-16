"""
DOCX -> Template IR.

Same graph, different addressing:

    Document > Section(heading) > {Paragraph label/value, Table}

Three fillable shapes are recognised:
  1. "Label:  ______"        -> inline value region inside one paragraph
  2. "Label:" then a blank    -> value region is the following paragraph
  3. 2-column key/value table -> right cell is the value region
  4. n-column table w/ header -> repeating collection (rows cloned from row 1)
"""

from __future__ import annotations

import re
from typing import Dict, List, Optional

from docx import Document
from docx.table import Table as DocxTable
from docx.text.paragraph import Paragraph

from ..ir import (ColumnIR, DocType, Location, Node, NodeType, Section,
                  StyleSignals, TableIR, TemplateIR, ValueFormat)
from .base import (BLANK_LINE_RE, TemplateParser, detect_value_format,
                   find_placeholder, looks_like_label, register)

LABEL_VALUE_RE = re.compile(r"^\s*([^:\n]{1,60}?)\s*:\s*(.*)$")


class DocxParser(TemplateParser):
    doc_type = DocType.DOCX

    def parse(self, path: str) -> TemplateIR:
        doc = Document(path)
        ir = TemplateIR(doc_type=DocType.DOCX, source_path=path,
                        meta={"paragraphs": len(doc.paragraphs), "tables": len(doc.tables)})

        current = Section(section_id="section::0", title="Document",
                          location=Location(DocType.DOCX, {"kind": "document"}))
        ir.sections.append(current)

        paras = doc.paragraphs
        for i, p in enumerate(paras):
            text = (p.text or "").strip()
            style = (p.style.name if p.style is not None else "") or ""
            sig = self._para_style(p)

            if not text:
                continue

            # ---- heading -> new section, and often also a fillable block
            if style.lower().startswith("heading") or style.lower() in ("title", "subtitle"):
                current = Section(section_id=f"section::{i}", title=text,
                                  location=Location(DocType.DOCX, {"kind": "paragraph",
                                                                   "para_index": i}))
                ir.sections.append(current)
                node = self._node(f"para::{i}", NodeType.SECTION_HEADER,
                                  {"kind": "paragraph", "para_index": i},
                                  text, current.title, sig)
                node.editable = False
                node.is_empty = False
                ir.nodes.append(node)
                current.node_ids.append(node.node_id)

                # A heading with a blank paragraph directly beneath it is
                # itself a label for the content that belongs under it -
                # "Summary" / "Decisions Taken" with nothing written yet.
                nxt = paras[i + 1] if i + 1 < len(paras) else None
                if nxt is not None and not (nxt.text or "").strip() and looks_like_label(text + ":"):
                    vnode = self._node(f"para::{i+1}", NodeType.VALUE_REGION,
                                       {"kind": "paragraph", "para_index": i + 1,
                                        "mode": "replace"},
                                       "", current.title, self._para_style(nxt))
                    vnode.label = text.rstrip(":").strip()
                    vnode.label_node_id = node.node_id
                    vnode.meta["multiline_capable"] = True
                    ir.nodes.append(vnode)
                    current.node_ids.append(vnode.node_id)
                continue

            ph = find_placeholder(text)
            if ph:
                node = self._node(f"para::{i}", NodeType.VALUE_REGION,
                                  {"kind": "paragraph", "para_index": i, "mode": "placeholder"},
                                  text, current.title, sig)
                node.placeholder = ph
                node.label = ph.replace("_", " ")
                node.is_empty = False
                ir.nodes.append(node)
                current.node_ids.append(node.node_id)
                continue

            m = LABEL_VALUE_RE.match(text)
            if m and looks_like_label(m.group(1) + ":"):
                label, tail = m.group(1).strip(), m.group(2).strip()
                tail_is_blank = (not tail) or bool(BLANK_LINE_RE.match(tail))
                if tail_is_blank:
                    node = self._node(f"para::{i}::value", NodeType.VALUE_REGION,
                                      {"kind": "paragraph", "para_index": i,
                                       "mode": "after_colon", "label_text": label},
                                      "", current.title, sig)
                    node.label = label
                    node.is_empty = True
                    node.meta["multiline_capable"] = False
                    ir.nodes.append(node)
                    current.node_ids.append(node.node_id)
                    continue

            # "Label:" alone with an empty paragraph beneath -> block value region
            if text.endswith(":") and looks_like_label(text):
                nxt = paras[i + 1] if i + 1 < len(paras) else None
                if nxt is not None and not (nxt.text or "").strip():
                    node = self._node(f"para::{i+1}", NodeType.VALUE_REGION,
                                      {"kind": "paragraph", "para_index": i + 1,
                                       "mode": "replace"},
                                      "", current.title, self._para_style(nxt))
                    node.label = text.rstrip(":").strip()
                    node.label_node_id = f"para::{i}"
                    node.meta["multiline_capable"] = True
                    ir.nodes.append(node)
                    current.node_ids.append(node.node_id)
                    continue

            static = self._node(f"para::{i}", NodeType.STATIC,
                                {"kind": "paragraph", "para_index": i},
                                text, current.title, sig)
            static.editable = False
            static.is_empty = False
            ir.nodes.append(static)
            current.node_ids.append(static.node_id)

        # ---- tables
        for ti, t in enumerate(doc.tables):
            self._parse_table(ir, ir.sections[-1], t, ti)

        return ir

    # ------------------------------------------------------------------
    def _parse_table(self, ir: TemplateIR, section: Section, t: DocxTable, ti: int) -> None:
        rows = t.rows
        if not rows:
            return
        header = [c.text.strip() for c in rows[0].cells]
        ncols = len(header)

        # 2-column key/value block: each row is a label/value pair, not a collection
        if ncols == 2 and self._is_key_value_table(t):
            for ri, row in enumerate(rows):
                key = row.cells[0].text.strip()
                val = row.cells[1].text.strip()
                if not key or not looks_like_label(key):
                    continue
                if val and not BLANK_LINE_RE.match(val) and not find_placeholder(val):
                    continue
                node = self._node(f"table::{ti}::r{ri}c1", NodeType.VALUE_REGION,
                                  {"kind": "table_cell", "table_index": ti, "row": ri, "col": 1},
                                  "", section.title, StyleSignals())
                node.label = key.rstrip(":*").strip()
                node.meta["multiline_capable"] = True
                ph = find_placeholder(val)
                if ph:
                    node.placeholder = ph
                ir.nodes.append(node)
                section.node_ids.append(node.node_id)
            return

        # otherwise: repeating collection
        if ncols < 2 or not any(header):
            return
        cols = [ColumnIR(index=i, header_text=h, location_hint={"col": i})
                for i, h in enumerate(header)]
        existing = sum(1 for r in rows[1:] if any(c.text.strip() for c in r.cells))
        table = TableIR(
            node_id=f"table::{ti}",
            location=Location(DocType.DOCX, {"kind": "table", "table_index": ti}),
            columns=cols, header_row=0, template_row=1,
            existing_data_rows=existing, section=section.title,
            meta={"n_rows": len(rows), "n_cols": ncols},
        )
        ir.tables.append(table)
        section.table_ids.append(table.node_id)

    def _is_key_value_table(self, t: DocxTable) -> bool:
        filled_left = sum(1 for r in t.rows if r.cells[0].text.strip())
        filled_right = sum(1 for r in t.rows if r.cells[1].text.strip())
        return filled_left >= max(2, len(t.rows) - 1) and filled_right <= max(1, len(t.rows) // 3)

    def _para_style(self, p: Paragraph) -> StyleSignals:
        run = p.runs[0] if p.runs else None
        f = run.font if run is not None else None
        return StyleSignals(
            bold=bool(run.bold) if run is not None else False,
            italic=bool(run.italic) if run is not None else False,
            font_size=float(f.size.pt) if f is not None and f.size else None,
            font_name=f.name if f is not None else None,
            alignment=str(p.alignment) if p.alignment is not None else None,
        )

    def _node(self, node_id: str, ntype: NodeType, parts: Dict, text: str,
              section: str, style: StyleSignals) -> Node:
        return Node(node_id=node_id, type=ntype,
                    location=Location(DocType.DOCX, parts),
                    text=text, section=section, style=style,
                    value_format=detect_value_format(None, text),
                    is_empty=not text.strip())


register("docx", DocxParser)
register("dotx", DocxParser)
